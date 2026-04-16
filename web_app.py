import math
import os
import re
from datetime import datetime
from io import BytesIO

import pandas as pd
from flask import Flask, jsonify, render_template, request, send_file
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

try:
    from reportlab.lib import colors as rl_colors
    from reportlab.lib.units import cm as rl_cm
    from reportlab.platypus import PageBreak, SimpleDocTemplate, Spacer, Table, TableStyle

    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False

try:
    from docx import Document as DocxDocument
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm as DocxCm
    from docx.shared import Pt

    DOCX_OK = True
except ImportError:
    DOCX_OK = False


app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024

PERIODOS = ["001", "002", "003", "FINAL"]
METRICAS_BASE = ["PROM", "PUESTO", "REPR", "ART", "CNT", "CPM", "EDF", "EMPIN"]
ALIAS_METRICAS = {
    "prom": "PROM",
    "promedio": "PROM",
    "puesto": "PUESTO",
    "repr": "REPR",
    "art": "ART",
    "cnt": "CNT",
    "cpm": "CPM",
    "edf": "EDF",
    "empi": "EMPIN",
    "empin": "EMPIN",
}
ALIAS_ESTUDIANTE = [
    "estudiante",
    "alumno",
    "alumna",
    "nombre",
    "nombres",
    "est",
    "id",
    "identificacion",
    "codigo",
    "documento",
]
ALIAS_PERIODO = ["periodo", "periodo academico", "periodo aca", "corte", "lapso"]
PATRONES_NO_ESTUDIANTE = [
    r"^\d+\.?\s*desempe[nñ]o",
    r"^desempe[nñ]o",
    r"consolidado",
    r"del curso",
    r"^grupo\b",
    r"totales?",
    r"resumen",
]
TAMANIO_PAGINAS = {
    "carta": (21.59, 27.94),
    "a4": (21.00, 29.70),
}


def normalizar_texto(valor: object) -> str:
    if pd.isna(valor):
        return ""
    texto = str(valor).strip()
    texto = re.sub(r"\s+", " ", texto)
    return texto


def clave_texto(valor: object) -> str:
    texto = normalizar_texto(valor).lower()
    texto = (
        texto.replace("á", "a")
        .replace("é", "e")
        .replace("í", "i")
        .replace("ó", "o")
        .replace("ú", "u")
    )
    return texto


def contiene_letras(valor: object) -> bool:
    texto = normalizar_texto(valor)
    return bool(re.search(r"[A-Za-zÁÉÍÓÚáéíóúÑñ]", texto))


def parece_numero(valor: object) -> bool:
    if pd.isna(valor):
        return False
    texto = normalizar_texto(valor).replace(",", ".")
    if not texto:
        return False
    return bool(re.fullmatch(r"[-+]?\d+(\.\d+)?", texto))


def es_columna_auxiliar(columna: object) -> bool:
    return bool(re.fullmatch(r"col_\d+(?:_\d+)?", clave_texto(columna)))


def es_etiqueta_estudiante_valida(texto: str) -> bool:
    if not texto:
        return False

    clave = clave_texto(texto)
    if len(clave) < 4:
        return False

    if not contiene_letras(texto):
        return False

    for patron in PATRONES_NO_ESTUDIANTE:
        if re.search(patron, clave):
            return False

    return True


def nombre_columna_unico(nombre: str, usados: set[str]) -> str:
    base = nombre or "COL"
    candidato = base
    sufijo = 2
    while candidato in usados:
        candidato = f"{base}_{sufijo}"
        sufijo += 1
    usados.add(candidato)
    return candidato


def detectar_fila_encabezado(df: pd.DataFrame) -> int:
    mejor_indice = 0
    mejor_puntaje = -1
    limite = min(len(df), 20)

    for indice in range(limite):
        puntaje = 0
        for valor in df.iloc[indice].tolist():
            clave = clave_texto(valor)
            if clave in ALIAS_METRICAS:
                puntaje += 4
            elif clave in ALIAS_PERIODO:
                puntaje += 3
            elif clave in ALIAS_ESTUDIANTE:
                puntaje += 3
            elif normalizar_texto(valor).upper() in PERIODOS:
                puntaje += 1
        if puntaje > mejor_puntaje:
            mejor_puntaje = puntaje
            mejor_indice = indice

    return mejor_indice


def preparar_dataframe(df_raw: pd.DataFrame) -> pd.DataFrame:
    df_limpio = df_raw.dropna(how="all").dropna(axis=1, how="all").reset_index(drop=True)
    if df_limpio.empty:
        raise ValueError("El archivo no contiene datos utilizables.")

    fila_encabezado = detectar_fila_encabezado(df_limpio)
    fila = df_limpio.iloc[fila_encabezado].tolist()
    usados: set[str] = set()
    encabezados: list[str] = []

    for indice, valor in enumerate(fila, start=1):
        texto = normalizar_texto(valor) or f"COL_{indice}"
        encabezados.append(nombre_columna_unico(texto, usados))

    df_datos = df_limpio.iloc[fila_encabezado + 1 :].copy()
    df_datos.columns = encabezados
    df_datos = df_datos.dropna(how="all").reset_index(drop=True)

    if df_datos.empty:
        raise ValueError("No se encontraron filas de datos debajo del encabezado detectado.")

    return df_datos


def detectar_columna_periodo(df: pd.DataFrame) -> str:
    mejor_columna = ""
    mejor_puntaje = -1.0

    for columna in df.columns:
        nombre = clave_texto(columna)
        valores = df[columna].apply(normalizar_texto)
        valores = valores[valores != ""]
        if valores.empty:
            continue

        coincidencias = valores.apply(lambda valor: valor.upper() in PERIODOS).sum()
        puntaje = coincidencias / len(valores)
        if nombre in ALIAS_PERIODO:
            puntaje += 1

        if puntaje > mejor_puntaje:
            mejor_puntaje = puntaje
            mejor_columna = columna

    if not mejor_columna or mejor_puntaje <= 0:
        raise ValueError("No fue posible identificar la columna del periodo.")
    return mejor_columna


def detectar_columnas_estudiante(df: pd.DataFrame, columna_periodo: str) -> tuple[str, str | None]:
    candidatos: list[tuple[float, str]] = []

    for columna in df.columns:
        if columna == columna_periodo:
            continue

        nombre = clave_texto(columna)
        serie = df[columna].apply(normalizar_texto)
        valores = serie[serie != ""]
        if valores.empty:
            continue

        longitud_media = valores.map(len).mean()
        variedad = valores.nunique() / max(len(valores), 1)
        puntaje = variedad + min(longitud_media / 20, 1)

        proporcion_letras = valores.apply(contiene_letras).mean()
        proporcion_numerica = valores.apply(parece_numero).mean()
        puntaje += proporcion_letras * 2
        puntaje -= proporcion_numerica * 2

        if any(alias in nombre for alias in ALIAS_ESTUDIANTE):
            puntaje += 3
        if valores.apply(lambda valor: valor.upper() in PERIODOS).mean() > 0.5:
            puntaje = -1

        candidatos.append((puntaje, columna))

    candidatos.sort(reverse=True)
    if not candidatos or candidatos[0][0] <= 0:
        raise ValueError("No fue posible identificar la columna del estudiante.")

    principal = candidatos[0][1]
    secundario = None
    for puntaje, columna in candidatos[1:]:
        nombre = clave_texto(columna)
        if puntaje > 0.5 and any(alias in nombre for alias in ["id", "codigo", "documento"]):
            secundario = columna
            break
    return principal, secundario


def detectar_columnas_metricas(df: pd.DataFrame, ignoradas: set[str]) -> dict[str, str]:
    mapeo: dict[str, str] = {}
    for columna in df.columns:
        if columna in ignoradas:
            continue

        clave = clave_texto(columna)
        if clave in ALIAS_METRICAS:
            mapeo[ALIAS_METRICAS[clave]] = columna
            continue

        serie = df[columna].apply(normalizar_texto)
        no_vacios = serie[serie != ""]
        if no_vacios.empty:
            continue

        proporcion_numerica = no_vacios.apply(parece_numero).mean()
        if proporcion_numerica >= 0.35:
            nombre_visible = normalizar_texto(columna).upper()
            if nombre_visible and not clave.startswith("col_"):
                mapeo[nombre_visible] = columna
    return mapeo


def fusionar_columnas_auxiliares(df: pd.DataFrame, columna_periodo: str) -> pd.DataFrame:
    trabajo = df.copy()
    columnas = list(trabajo.columns)
    a_eliminar: list[str] = []

    periodos = trabajo[columna_periodo].apply(lambda valor: normalizar_texto(valor).upper())
    es_final = periodos == "FINAL"

    for indice, columna in enumerate(columnas):
        if indice == 0 or not es_columna_auxiliar(columna):
            continue

        destino = columnas[indice - 1]
        if es_columna_auxiliar(destino):
            continue

        destino_vacio = trabajo[destino].apply(lambda v: pd.isna(v) or normalizar_texto(v) == "")
        auxiliar_con_dato = ~trabajo[columna].apply(lambda v: pd.isna(v) or normalizar_texto(v) == "")

        mascara_final = es_final & destino_vacio & auxiliar_con_dato
        if mascara_final.any():
            trabajo.loc[mascara_final, destino] = trabajo.loc[mascara_final, columna]

        destino_vacio = trabajo[destino].apply(lambda v: pd.isna(v) or normalizar_texto(v) == "")
        mascara_general = destino_vacio & auxiliar_con_dato
        if mascara_general.any():
            trabajo.loc[mascara_general, destino] = trabajo.loc[mascara_general, columna]

        a_eliminar.append(columna)

    if a_eliminar:
        trabajo = trabajo.drop(columns=a_eliminar, errors="ignore")
    return trabajo


def construir_etiqueta_estudiante(
    fila: pd.Series,
    columna_principal: str,
    columna_secundaria: str | None,
) -> str:
    principal = normalizar_texto(fila.get(columna_principal, ""))
    secundario = normalizar_texto(fila.get(columna_secundaria, "")) if columna_secundaria else ""
    if principal and secundario and principal != secundario:
        return f"{principal} ({secundario})"
    return principal or secundario


def normalizar_tabla_fuente(df: pd.DataFrame) -> tuple[pd.DataFrame, list[str]]:
    columna_periodo = detectar_columna_periodo(df)
    df = fusionar_columnas_auxiliares(df, columna_periodo)

    columna_periodo = detectar_columna_periodo(df)
    columna_estudiante, columna_id = detectar_columnas_estudiante(df, columna_periodo)

    ignoradas = {columna_periodo, columna_estudiante}
    if columna_id:
        ignoradas.add(columna_id)

    columnas_metricas = detectar_columnas_metricas(df, ignoradas)
    if not columnas_metricas:
        raise ValueError("No se detectaron columnas de metricas en el archivo.")

    metricas_ordenadas: list[str] = []
    for metrica in METRICAS_BASE:
        if metrica in columnas_metricas:
            metricas_ordenadas.append(metrica)
    for metrica in columnas_metricas.keys():
        if metrica not in metricas_ordenadas:
            metricas_ordenadas.append(metrica)

    trabajo = df.copy()
    trabajo[columna_estudiante] = trabajo[columna_estudiante].ffill()
    if columna_id:
        trabajo[columna_id] = trabajo[columna_id].ffill()

    trabajo["_periodo"] = trabajo[columna_periodo].apply(lambda v: normalizar_texto(v).upper())
    trabajo = trabajo[trabajo["_periodo"].isin(PERIODOS)].copy()
    if trabajo.empty:
        raise ValueError("No se encontraron filas con periodos 001, 002, 003 o FINAL.")

    trabajo["_estudiante"] = trabajo.apply(
        lambda fila: construir_etiqueta_estudiante(fila, columna_estudiante, columna_id),
        axis=1,
    )
    trabajo = trabajo[trabajo["_estudiante"].apply(es_etiqueta_estudiante_valida)].copy()
    if trabajo.empty:
        raise ValueError("No se detectaron estudiantes validos en el archivo.")

    columnas_finales = ["_estudiante", "_periodo"]
    renombrar: dict[str, str] = {}

    for metrica, columna_original in columnas_metricas.items():
        columnas_finales.append(columna_original)
        renombrar[columna_original] = metrica

    normalizado = trabajo[columnas_finales].rename(columns=renombrar).copy()
    for metrica in metricas_ordenadas:
        if metrica in normalizado.columns:
            normalizado[metrica] = pd.to_numeric(normalizado[metrica], errors="coerce")

    return normalizado, metricas_ordenadas


def formatear_valor(valor: object) -> object:
    if pd.isna(valor):
        return ""
    if isinstance(valor, float):
        if math.isclose(valor, round(valor)):
            return int(round(valor))
        return round(valor, 2)
    return valor


def crear_matriz_estudiante(df_estudiante: pd.DataFrame, metricas: list[str], periodos: list[str]) -> pd.DataFrame:
    matriz = pd.DataFrame(index=metricas, columns=periodos, dtype=object).fillna("")

    for _, fila in df_estudiante.iterrows():
        periodo = fila["_periodo"]
        if periodo not in periodos:
            continue
        for metrica in metricas:
            if metrica in df_estudiante.columns:
                valor = fila.get(metrica)
                if pd.notna(valor):
                    matriz.loc[metrica, periodo] = formatear_valor(valor)
    return matriz


def etiqueta_nivelacion(promedio: float | None) -> str:
    if promedio is None or pd.isna(promedio):
        return "SIN NOTA"
    if promedio >= 4.6:
        return "SUPERIOR"
    if promedio >= 4.0:
        return "ALTO"
    if promedio >= 3.0:
        return "BASICO"
    return "BAJO"


def aplicar_logica_nivelacion(matriz: pd.DataFrame, periodos: list[str]) -> pd.DataFrame:
    resultado = matriz.copy()
    if "PROM" not in resultado.index:
        return resultado

    niveles: dict[str, str] = {}
    requiere: dict[str, str] = {}
    for periodo in periodos:
        valor = resultado.loc["PROM", periodo] if periodo in resultado.columns else ""
        try:
            promedio = float(valor)
        except (TypeError, ValueError):
            promedio = None

        nivel = etiqueta_nivelacion(promedio)
        niveles[periodo] = nivel
        requiere[periodo] = "SI" if nivel == "BAJO" else "NO"

    resultado.loc["NIVELACION"] = pd.Series(niveles)
    resultado.loc["REQUIERE_NIVELACION"] = pd.Series(requiere)
    return resultado


def obtener_configuracion_pagina(estudiantes_por_fila: int, filas_por_hoja: int) -> tuple[int, int]:
    if estudiantes_por_fila <= 0 or filas_por_hoja <= 0:
        raise ValueError("Estudiantes por fila y filas por hoja deben ser mayores a cero.")
    return estudiantes_por_fila, filas_por_hoja


def es_orientacion_horizontal(orientacion: str) -> bool:
    return orientacion == "horizontal"


def obtener_tamano_pagina_orientado(
    tamano_pagina: str,
    pagina_ancho_cm: float,
    pagina_alto_cm: float,
    orientacion: str,
) -> tuple[float, float]:
    if tamano_pagina in TAMANIO_PAGINAS:
        ancho_cm, alto_cm = TAMANIO_PAGINAS[tamano_pagina]
    else:
        ancho_cm, alto_cm = pagina_ancho_cm, pagina_alto_cm

    if es_orientacion_horizontal(orientacion):
        return max(ancho_cm, alto_cm), min(ancho_cm, alto_cm)
    return min(ancho_cm, alto_cm), max(ancho_cm, alto_cm)


def configurar_hoja_excel(hoja, tamano_pagina: str, orientacion: str) -> None:
    if tamano_pagina == "carta":
        hoja.page_setup.paperSize = hoja.PAPERSIZE_LETTER
    else:
        hoja.page_setup.paperSize = hoja.PAPERSIZE_A4

    hoja.page_setup.orientation = "landscape" if es_orientacion_horizontal(orientacion) else "portrait"
    hoja.page_setup.fitToWidth = 1
    hoja.page_setup.fitToHeight = 1
    hoja.sheet_view.showGridLines = False
    hoja.page_margins.left = 0.2
    hoja.page_margins.right = 0.2
    hoja.page_margins.top = 0.3
    hoja.page_margins.bottom = 0.3


def escribir_bloque_estudiante_excel(
    hoja,
    fila_inicio: int,
    columna_inicio: int,
    estudiante: str,
    matriz: pd.DataFrame,
    metricas: list[str],
    periodos: list[str],
) -> None:
    relleno_titulo = PatternFill("solid", fgColor="D9D9D9")
    relleno_encabezado = PatternFill("solid", fgColor="EDEDED")
    borde = Border(
        left=Side(style="thin", color="808080"),
        right=Side(style="thin", color="808080"),
        top=Side(style="thin", color="808080"),
        bottom=Side(style="thin", color="808080"),
    )
    centrado = Alignment(horizontal="center", vertical="center", wrap_text=True)
    fuente_titulo = Font(name="Calibri", size=10, bold=True)
    fuente_normal = Font(name="Calibri", size=9)
    ultima_columna = columna_inicio + len(periodos)

    hoja.merge_cells(
        start_row=fila_inicio,
        start_column=columna_inicio,
        end_row=fila_inicio,
        end_column=ultima_columna,
    )
    celda_titulo = hoja.cell(row=fila_inicio, column=columna_inicio, value=estudiante)
    celda_titulo.fill = relleno_titulo
    celda_titulo.border = borde
    celda_titulo.alignment = centrado
    celda_titulo.font = fuente_titulo

    for columna_offset, periodo in enumerate(["", *periodos]):
        celda = hoja.cell(row=fila_inicio + 1, column=columna_inicio + columna_offset, value=periodo)
        celda.fill = relleno_encabezado
        celda.border = borde
        celda.alignment = centrado
        celda.font = fuente_titulo

    for fila_offset, metrica in enumerate(metricas, start=2):
        celda_metrica = hoja.cell(row=fila_inicio + fila_offset, column=columna_inicio, value=metrica)
        celda_metrica.fill = relleno_encabezado
        celda_metrica.border = borde
        celda_metrica.alignment = centrado
        celda_metrica.font = fuente_titulo

        for columna_offset, periodo in enumerate(periodos, start=1):
            valor = matriz.loc[metrica, periodo] if metrica in matriz.index and periodo in matriz.columns else ""
            celda_valor = hoja.cell(
                row=fila_inicio + fila_offset,
                column=columna_inicio + columna_offset,
                value=valor,
            )
            celda_valor.border = borde
            celda_valor.alignment = centrado
            celda_valor.font = fuente_normal

    for columna in range(columna_inicio, ultima_columna + 1):
        hoja.cell(row=fila_inicio, column=columna).border = borde


def ajustar_anchos_hoja_excel(hoja, columnas_por_bloque: int) -> None:
    for columna in range(1, hoja.max_column + 1):
        letra = get_column_letter(columna)
        maximo = 0
        for celda in hoja[letra]:
            valor = "" if celda.value is None else str(celda.value)
            maximo = max(maximo, len(valor))

        if (columna - 1) % columnas_por_bloque == 0:
            hoja.column_dimensions[letra].width = min(max(maximo + 2, 10), 18)
        else:
            hoja.column_dimensions[letra].width = min(max(maximo + 2, 7), 12)


def construir_excel_reporte(
    reportes: dict[str, pd.DataFrame],
    periodos: list[str],
    metricas: list[str],
    estudiantes_por_fila: int,
    filas_por_hoja: int,
    tamano_pagina: str,
    orientacion: str,
) -> BytesIO:
    libro = Workbook()
    libro.remove(libro.active)

    estudiantes_por_fila, filas_por_hoja = obtener_configuracion_pagina(estudiantes_por_fila, filas_por_hoja)
    columnas_por_bloque = 1 + len(periodos)

    alto_contenido = 2 + len(metricas)
    alto_separacion = 1
    alto_bloque = alto_contenido + alto_separacion

    estudiantes = list(reportes.items())
    por_hoja = estudiantes_por_fila * filas_por_hoja

    for inicio in range(0, len(estudiantes), por_hoja):
        hoja = libro.create_sheet(title=f"Pagina_{inicio // por_hoja + 1}")
        configurar_hoja_excel(hoja, tamano_pagina, orientacion)

        for indice_bloque, (estudiante, matriz) in enumerate(estudiantes[inicio : inicio + por_hoja]):
            fila_bloque = indice_bloque // estudiantes_por_fila
            columna_bloque = indice_bloque % estudiantes_por_fila
            fila_inicio = 1 + fila_bloque * alto_bloque
            columna_inicio = 1 + columna_bloque * columnas_por_bloque

            escribir_bloque_estudiante_excel(
                hoja,
                fila_inicio,
                columna_inicio,
                estudiante,
                matriz,
                metricas,
                periodos,
            )

            hoja.row_dimensions[fila_inicio].height = 18
            hoja.row_dimensions[fila_inicio + 1].height = 14
            for fila_offset in range(2, alto_contenido):
                hoja.row_dimensions[fila_inicio + fila_offset].height = 13
            hoja.row_dimensions[fila_inicio + alto_contenido].height = 8

        ajustar_anchos_hoja_excel(hoja, columnas_por_bloque)

    salida = BytesIO()
    libro.save(salida)
    salida.seek(0)
    return salida


def construir_pdf_reporte(
    reportes: dict[str, pd.DataFrame],
    periodos: list[str],
    metricas: list[str],
    estudiantes_por_fila: int,
    filas_por_hoja: int,
    tamano_pagina: str,
    orientacion: str,
    pagina_ancho_cm: float,
    pagina_alto_cm: float,
) -> BytesIO:
    if not REPORTLAB_OK:
        raise ValueError("Para exportar a PDF instala reportlab (pip install reportlab).")

    ancho_cm, alto_cm = obtener_tamano_pagina_orientado(
        tamano_pagina,
        pagina_ancho_cm,
        pagina_alto_cm,
        orientacion,
    )
    ancho_pt = ancho_cm * rl_cm
    alto_pt = alto_cm * rl_cm

    estudiantes_por_fila, filas_por_hoja = obtener_configuracion_pagina(estudiantes_por_fila, filas_por_hoja)
    por_hoja = estudiantes_por_fila * filas_por_hoja

    num_periodos = len(periodos)
    margen_pt = 0.5 * rl_cm
    disponible = ancho_pt - 2 * margen_pt
    gap_pt = 3
    bloque_pt = (disponible - gap_pt * (estudiantes_por_fila - 1)) / estudiantes_por_fila
    col_label_pt = bloque_pt * 0.30
    col_per_pt = (bloque_pt - col_label_pt) / num_periodos
    col_anchos = [col_label_pt] + [col_per_pt] * num_periodos
    altos = [14, 11] + [10] * len(metricas)

    estilo_inner = TableStyle(
        [
            ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("GRID", (0, 0), (-1, -1), 0.4, rl_colors.grey),
            ("BACKGROUND", (0, 0), (-1, 0), rl_colors.HexColor("#D9D9D9")),
            ("BACKGROUND", (0, 1), (-1, 1), rl_colors.HexColor("#EDEDED")),
            ("BACKGROUND", (0, 2), (0, -1), rl_colors.HexColor("#EDEDED")),
            ("FONTNAME", (0, 0), (-1, 1), "Helvetica-Bold"),
            ("FONTNAME", (0, 2), (0, -1), "Helvetica-Bold"),
            ("SPAN", (0, 0), (-1, 0)),
        ]
    )

    def tabla_estudiante(nombre_estudiante: str, matriz_estudiante: pd.DataFrame):
        encabezado = [nombre_estudiante] + [""] * num_periodos
        sub = [""] + periodos
        datos = [encabezado, sub]
        for metrica in metricas:
            fila = [metrica]
            for periodo in periodos:
                valor = ""
                if metrica in matriz_estudiante.index and periodo in matriz_estudiante.columns:
                    raw = matriz_estudiante.loc[metrica, periodo]
                    valor = "" if raw == "" else str(raw)
                fila.append(valor)
            datos.append(fila)

        tabla = Table(datos, colWidths=col_anchos, rowHeights=altos)
        tabla.setStyle(estilo_inner)
        return tabla

    story = []
    todos = list(reportes.items())
    primera_pagina = True

    for inicio in range(0, len(todos), por_hoja):
        if not primera_pagina:
            story.append(PageBreak())
        primera_pagina = False

        grupo = todos[inicio : inicio + por_hoja]
        for fila_indice in range(0, len(grupo), estudiantes_por_fila):
            fila_ests = grupo[fila_indice : fila_indice + estudiantes_por_fila]
            tablas = [tabla_estudiante(nombre, matriz) for nombre, matriz in fila_ests]

            while len(tablas) < estudiantes_por_fila:
                n_filas = 2 + len(metricas)
                t_vacio = Table([[""] * (1 + num_periodos)] * n_filas, colWidths=col_anchos, rowHeights=altos)
                tablas.append(t_vacio)

            maestra = Table([tablas], colWidths=[bloque_pt] * estudiantes_por_fila, hAlign="LEFT")
            maestra.setStyle(
                TableStyle(
                    [
                        ("LEFTPADDING", (0, 0), (-1, -1), 0),
                        ("RIGHTPADDING", (0, 0), (-1, -1), gap_pt),
                        ("TOPPADDING", (0, 0), (-1, -1), 0),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ]
                )
            )
            story.append(maestra)
            story.append(Spacer(1, 6))

    salida = BytesIO()
    SimpleDocTemplate(
        salida,
        pagesize=(ancho_pt, alto_pt),
        leftMargin=margen_pt,
        rightMargin=margen_pt,
        topMargin=margen_pt,
        bottomMargin=margen_pt,
    ).build(story)
    salida.seek(0)
    return salida


def construir_word_reporte(
    reportes: dict[str, pd.DataFrame],
    periodos: list[str],
    metricas: list[str],
    estudiantes_por_fila: int,
    filas_por_hoja: int,
    tamano_pagina: str,
    orientacion: str,
    pagina_ancho_cm: float,
    pagina_alto_cm: float,
) -> BytesIO:
    if not DOCX_OK:
        raise ValueError("Para exportar a Word instala python-docx (pip install python-docx).")

    ancho_cm, alto_cm = obtener_tamano_pagina_orientado(
        tamano_pagina,
        pagina_ancho_cm,
        pagina_alto_cm,
        orientacion,
    )

    estudiantes_por_fila, filas_por_hoja = obtener_configuracion_pagina(estudiantes_por_fila, filas_por_hoja)
    por_hoja = estudiantes_por_fila * filas_por_hoja
    cols_por_est = 1 + len(periodos)

    documento = DocxDocument()
    seccion = documento.sections[0]
    seccion.orientation = WD_ORIENT.LANDSCAPE if es_orientacion_horizontal(orientacion) else WD_ORIENT.PORTRAIT
    seccion.page_width = DocxCm(ancho_cm)
    seccion.page_height = DocxCm(alto_cm)
    seccion.left_margin = DocxCm(0.8)
    seccion.right_margin = DocxCm(0.8)
    seccion.top_margin = DocxCm(0.8)
    seccion.bottom_margin = DocxCm(0.8)

    todos = list(reportes.items())
    primera_pagina = True

    for inicio in range(0, len(todos), por_hoja):
        if not primera_pagina:
            documento.add_page_break()
        primera_pagina = False

        grupo = todos[inicio : inicio + por_hoja]
        num_filas_tabla = 2 + len(metricas)

        for fila_indice in range(0, len(grupo), estudiantes_por_fila):
            fila_ests = grupo[fila_indice : fila_indice + estudiantes_por_fila]
            cantidad_est = len(fila_ests)
            total_cols = cantidad_est * cols_por_est

            tabla = documento.add_table(rows=num_filas_tabla, cols=total_cols)
            tabla.style = "Table Grid"

            ancho_disp = ancho_cm - 1.6
            ancho_label = (ancho_disp / cantidad_est) * 0.30
            ancho_periodo = ((ancho_disp / cantidad_est) - ancho_label) / max(len(periodos), 1)

            for col_est, (nombre_est, matriz) in enumerate(fila_ests):
                col_base = col_est * cols_por_est

                celda_titulo = tabla.cell(0, col_base)
                for col in range(1, cols_por_est):
                    celda_titulo = celda_titulo.merge(tabla.cell(0, col_base + col))

                p_titulo = celda_titulo.paragraphs[0]
                p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_titulo = p_titulo.add_run(str(nombre_est).upper())
                run_titulo.bold = True
                run_titulo.font.size = Pt(7)

                for offset, texto in enumerate([""] + periodos):
                    p = tabla.cell(1, col_base + offset).paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run(texto)
                    r.bold = True
                    r.font.size = Pt(7)

                for fila_offset, metrica in enumerate(metricas, start=2):
                    p_met = tabla.cell(fila_offset, col_base).paragraphs[0]
                    p_met.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r_met = p_met.add_run(metrica)
                    r_met.bold = True
                    r_met.font.size = Pt(7)

                    for col_offset, periodo in enumerate(periodos, start=1):
                        valor = ""
                        if metrica in matriz.index and periodo in matriz.columns:
                            raw = matriz.loc[metrica, periodo]
                            valor = "" if raw == "" else str(raw)

                        p_val = tabla.cell(fila_offset, col_base + col_offset).paragraphs[0]
                        p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r_val = p_val.add_run(valor)
                        r_val.font.size = Pt(7)

            for col_est in range(cantidad_est):
                col_base = col_est * cols_por_est
                for row in tabla.rows:
                    for col_offset in range(cols_por_est):
                        celda = row.cells[col_base + col_offset]
                        celda.width = DocxCm(ancho_label if col_offset == 0 else ancho_periodo)

            documento.add_paragraph("")

    salida = BytesIO()
    documento.save(salida)
    salida.seek(0)
    return salida


def generar_reportes_desde_excel(archivo_excel, periodos: list[str]) -> tuple[dict[str, pd.DataFrame], list[str]]:
    if not periodos:
        raise ValueError("Debes seleccionar al menos un periodo.")

    df_raw = pd.read_excel(archivo_excel, header=None, dtype=object)
    df_datos = preparar_dataframe(df_raw)
    df_normalizado, metricas_base = normalizar_tabla_fuente(df_datos)

    reportes: dict[str, pd.DataFrame] = {}
    for estudiante, grupo in df_normalizado.groupby("_estudiante", sort=True):
        matriz = crear_matriz_estudiante(grupo, metricas_base, periodos)
        reportes[estudiante] = matriz

    if not reportes:
        raise ValueError("No se generaron reportes individuales.")

    return reportes, metricas_base


def config_por_defecto() -> dict[str, object]:
    return {
        "formato": "excel",
        "estudiantes_por_fila": 7,
        "filas_por_hoja": 3,
        "tamano_pagina": "carta",
        "orientacion": "horizontal",
        "pagina_ancho_cm": 21.59,
        "pagina_alto_cm": 27.94,
    }


def leer_config_formulario(formulario) -> dict[str, object]:
    config = config_por_defecto()

    config["formato"] = formulario.get("formato", str(config["formato"]))
    config["tamano_pagina"] = formulario.get("tamano_pagina", str(config["tamano_pagina"]))
    config["orientacion"] = formulario.get("orientacion", str(config["orientacion"]))

    try:
        config["estudiantes_por_fila"] = int(formulario.get("estudiantes_por_fila", config["estudiantes_por_fila"]))
        config["filas_por_hoja"] = int(formulario.get("filas_por_hoja", config["filas_por_hoja"]))
        config["pagina_ancho_cm"] = float(formulario.get("pagina_ancho_cm", config["pagina_ancho_cm"]))
        config["pagina_alto_cm"] = float(formulario.get("pagina_alto_cm", config["pagina_alto_cm"]))
    except ValueError as error:
        raise ValueError(f"Configuracion numerica invalida: {error}") from error

    if config["formato"] not in {"excel", "pdf", "word"}:
        raise ValueError("Formato de salida invalido.")
    if config["tamano_pagina"] not in {"carta", "a4", "custom"}:
        raise ValueError("Tamano de pagina invalido.")
    if config["orientacion"] not in {"vertical", "horizontal"}:
        raise ValueError("Orientacion invalida.")

    obtener_configuracion_pagina(config["estudiantes_por_fila"], config["filas_por_hoja"])

    if config["tamano_pagina"] == "custom":
        if config["pagina_ancho_cm"] <= 0 or config["pagina_alto_cm"] <= 0:
            raise ValueError("Las dimensiones personalizadas deben ser mayores a cero.")

    return config


def render_index_error(error: str, periodos_default: list[str], config: dict[str, object], status: int = 400):
    return (
        render_template(
            "index.html",
            periodos=PERIODOS,
            periodos_default=periodos_default or ["001"],
            config=config,
            error=error,
        ),
        status,
    )


@app.route("/health", methods=["GET"])
def health():
    return {"status": "ok"}, 200


@app.route("/", methods=["GET"])
def index():
    return render_template(
        "index.html",
        periodos=PERIODOS,
        periodos_default=["001"],
        config=config_por_defecto(),
        error=None,
    )


@app.route("/generar", methods=["POST"])
def generar():
    archivo = request.files.get("archivo_excel")
    periodos = request.form.getlist("periodos")

    try:
        config = leer_config_formulario(request.form)
    except ValueError as error:
        return render_index_error(str(error), periodos, config_por_defecto())

    if archivo is None or archivo.filename == "":
        return render_index_error("Debes cargar un archivo Excel (.xlsx).", periodos, config)

    if not archivo.filename.lower().endswith(".xlsx"):
        return render_index_error("Formato invalido. Solo se permite archivo .xlsx.", periodos, config)

    if config["formato"] == "pdf" and not REPORTLAB_OK:
        return render_index_error("Falta dependencia reportlab para exportar PDF.", periodos, config)

    if config["formato"] == "word" and not DOCX_OK:
        return render_index_error("Falta dependencia python-docx para exportar Word.", periodos, config)

    try:
        archivo.stream.seek(0)
        reportes, metricas = generar_reportes_desde_excel(archivo, periodos)

        if config["formato"] == "excel":
            stream = construir_excel_reporte(
                reportes=reportes,
                periodos=periodos,
                metricas=metricas,
                estudiantes_por_fila=config["estudiantes_por_fila"],
                filas_por_hoja=config["filas_por_hoja"],
                tamano_pagina=config["tamano_pagina"],
                orientacion=config["orientacion"],
            )
            extension = "xlsx"
            mimetype = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        elif config["formato"] == "pdf":
            stream = construir_pdf_reporte(
                reportes=reportes,
                periodos=periodos,
                metricas=metricas,
                estudiantes_por_fila=config["estudiantes_por_fila"],
                filas_por_hoja=config["filas_por_hoja"],
                tamano_pagina=config["tamano_pagina"],
                orientacion=config["orientacion"],
                pagina_ancho_cm=config["pagina_ancho_cm"],
                pagina_alto_cm=config["pagina_alto_cm"],
            )
            extension = "pdf"
            mimetype = "application/pdf"
        else:
            stream = construir_word_reporte(
                reportes=reportes,
                periodos=periodos,
                metricas=metricas,
                estudiantes_por_fila=config["estudiantes_por_fila"],
                filas_por_hoja=config["filas_por_hoja"],
                tamano_pagina=config["tamano_pagina"],
                orientacion=config["orientacion"],
                pagina_ancho_cm=config["pagina_ancho_cm"],
                pagina_alto_cm=config["pagina_alto_cm"],
            )
            extension = "docx"
            mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    except Exception as error:
        return render_index_error(f"No fue posible generar el reporte: {error}", periodos, config)

    marca_tiempo = datetime.now().strftime("%Y%m%d_%H%M%S")
    nombre = f"reporte_boletines_{marca_tiempo}.{extension}"
    return send_file(stream, as_attachment=True, download_name=nombre, mimetype=mimetype)


@app.errorhandler(413)
def archivo_muy_grande(_error):
    return jsonify({"error": "El archivo supera el tamano permitido de 16 MB"}), 413


if __name__ == "__main__":
    puerto = int(os.environ.get("PORT", "5000"))
    debug = os.environ.get("FLASK_DEBUG", "0") == "1"
    app.run(debug=debug, host="0.0.0.0", port=puerto)
